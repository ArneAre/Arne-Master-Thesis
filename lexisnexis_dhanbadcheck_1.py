# -*- coding: utf-8 -*-
"""
Created on Fri Apr 19 18:12:09 2024

@author: arnea
"""

import pandas as pd
from docx import Document

# Load the DataFrame
df = pd.read_csv('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/df_cleanedv5.csv')  # Update the path to your CSV file

# Filter to include only articles containing the word 'Dhanbad' (case insensitive)
df_filtered = df[df['Article_Text'].str.contains('Dhanbad', case=False, na=False)]

# Create a new Word document
doc = Document()
doc.add_heading('Articles Containing the Word "Dhanbad"', level=1)

# Add articles to the document
for index, row in df_filtered.iterrows():
    doc.add_heading(f'Title: {row["Title"]}', level=2)
    doc.add_paragraph(f'Date: {row["Date"]}')
    doc.add_paragraph(f'Publication: {row["Publication"]}')
    doc.add_paragraph('Article Text:')
    doc.add_paragraph(row['Article_Text'])
    doc.add_page_break()  # Add a page break after each article

# Save the document
doc.save('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/Articles_with_Dhanbadv2.docx')  # Update the path to where you want to save the document
