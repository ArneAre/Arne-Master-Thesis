# -*- coding: utf-8 -*-
"""
Created on Wed Apr 17 17:26:03 2024

@author: arnea
"""

import pandas as pd
from docx import Document
from nltk.tokenize import word_tokenize
import nltk
nltk.download('punkt')

# Load the DataFrame
df = pd.read_csv('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/df_no_duplicatesv3.csv')

# Ensuring the Article_Text column is a string
df['Article_Text'] = df['Article_Text'].astype(str)

# Calculate word counts
df['WordCount'] = df['Article_Text'].apply(lambda x: len(word_tokenize(x)))

# Filter articles with 70 to 99 words
filtered_df = df[(df['WordCount'] >= 70) & (df['WordCount'] <= 99)]

# Create a new Word document
doc = Document()

# Add a heading
doc.add_heading('Articles with 70 to 99 Words', level=1)

# Iterate over filtered articles
for index, row in filtered_df.iterrows():
    doc.add_heading(f"Article {index + 1}: {row['Title']}", level=2)
    doc.add_paragraph(f"Date: {row['Date']}")
    doc.add_paragraph(f"Publication: {row['Publication'] if 'Publication' in df.columns else 'Not available'}")
    doc.add_paragraph("Article Text:")
    doc.add_paragraph(row['Article_Text'])
    doc.add_paragraph("--------")  # Add a line for separation between articles

# Save the document
doc_path = 'C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/Articles_70_to_99_words.docx'
doc.save(doc_path)
print(f"Document saved to {doc_path}")
