# -*- coding: utf-8 -*-
"""
Created on Wed Apr 17 15:39:17 2024

@author: arnea
"""

import pandas as pd
from docx import Document
from docx.shared import Pt
from datetime import datetime

# Load my DataFrame
df = pd.read_csv('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/df_preprocessedv4.csv')
df['Date'] = pd.to_datetime(df['Date'])

# Create a new Word document
doc = Document()

# Organizing data by publication and then by Date
grouped = df.groupby('Publication')

# Loop through each publication
for publication, data in grouped:
    doc.add_heading(publication, level=1)
    # Sample 30 articles evenly across time slices
    time_slices = pd.qcut(data['Date'], q=30, duplicates='drop')
    sampled_data = data.groupby(time_slices).apply(lambda x: x.sample(n=1 if len(x) > 0 else 0)).reset_index(drop=True)

    # Add articles to the document
    for _, row in sampled_data.iterrows():
        title = row['Title']
        date = row['Date'].strftime('%Y-%m-%d')
        text = row['Article_Text']

        doc.add_heading(f'Title: {title}', level=2)
        doc.add_paragraph(f'Date: {date}')
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)  # Set the font size

        doc.add_page_break()

# Save the document
doc.save('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/articles_checkv1.docx')
