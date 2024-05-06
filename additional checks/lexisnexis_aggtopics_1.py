# -*- coding: utf-8 -*-
"""
Created on Thu Apr 11 14:19:08 2024

@author: arnea
"""

import pandas as pd
from docx import Document

topics_df = pd.read_csv("C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/dtm_topics_over_timev6.csv")

# Renaming mislabeled columns
topics_df.rename(columns={'Words': 'Weights', 'Weights': 'Words'}, inplace=True)

# Aggregate words for each topic across all time slices
aggregated_topics = topics_df.groupby('TopicNum').agg({
    'Words': lambda words: ' '.join([word for sublist in words for word in sublist])
}).reset_index()

# Renaming columns for clarity
aggregated_topics.columns = ['Topic Number', 'Aggregated Words']

# Display the DataFrame
print(aggregated_topics)

aggregated_topics.to_csv('aggregated_topicsv5.csv', index=False)

# Create a Word document
doc = Document()
doc.add_heading('Aggregated Topics', level=1)

# Add aggregated topics to the Word document
for index, row in aggregated_topics.iterrows():
    doc.add_heading(f'Topic {row["Topic Number"]}', level=2)
    doc.add_paragraph(row['Aggregated Words'])

# Save the document
doc.save('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/aggregated_topics6.docx')
