# -*- coding: utf-8 -*-
"""
Created on Thu Apr 18 15:00:34 2024

@author: arnea
"""

import pandas as pd
from docx import Document

# Load the DataFrame
df = pd.read_csv('C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/df_preprocessedv6.csv')

# Ensure the Article_Text column is a string
df['Article_Text'] = df['Article_Text'].astype(str)

# Filter articles containing any of the keywords within larger strings
keywords = ['html', 'www']
pattern = '|'.join(keywords)  # Adjust pattern to match within larger strings
df['Contains_Keywords'] = df['Article_Text'].str.contains(pattern, case=False, na=False)

# Filter to get only rows where Contains_Keywords is True
filtered_df = df[df['Contains_Keywords']]

# Create a new Word document
doc = Document()

# Add a heading
doc.add_heading('Articles Containing Keywords: html, www', level=1)

# Iterate over filtered articles
for index, row in filtered_df.iterrows():
    doc.add_heading(f"Article {index + 1}: {row['Title']}", level=2)
    doc.add_paragraph(f"Date: {row['Date']}")
    doc.add_paragraph(f"Publication: {row['Publication'] if 'Publication' in df.columns else 'Not available'}")
    doc.add_paragraph("Article Text:")
    doc.add_paragraph(row['Article_Text'])
    doc.add_paragraph("--------")  # Add a line for separation between articles

# Save the document
doc_path = 'C:/Users/arnea/OneDrive/Desktop/Thesis/Work/Python/Articles_with_Keywords2.docx'
doc.save(doc_path)
print(f"Document saved to {doc_path}")
